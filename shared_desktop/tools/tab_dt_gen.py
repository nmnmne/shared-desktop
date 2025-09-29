from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import base64
from collections import defaultdict

class GenerateDetectorTable(APIView):
    def post(self, request):
        detectors = request.data.get('detectors', [])
        
        if not detectors:
            return Response({"error": "Не указаны детекторы"}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            expanded_detectors = self._expand_detectors(detectors)
            
            # Подготовка данных для расчета КИ ПД-16
            ki_pd16_map = self._calculate_ki_pd16_map(expanded_detectors)
            
            # Генерируем данные для таблицы
            table_data = []
            detector_idx = 1
            
            for detector_name in expanded_detectors:
                detector_data = self._generate_detector_data(detector_name, detector_idx, expanded_detectors, ki_pd16_map)
                table_data.append(detector_data)
                detector_idx += 1
            
            # Заменяем все "7" на "6" в столбце offset
            for row in table_data:
                if row.get('offset') == '7':
                    row['offset'] = '6'

            # Заменяем все "3" на "2" в столбце offset
            for row in table_data:
                if row.get('offset') == '3':
                    row['offset'] = '2'

            # Создаем DOCX документ
            document = self._create_document(table_data)
            
            # Сохраняем документ в BytesIO
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            
            # Кодируем файл в base64 для включения в JSON
            file_base64 = base64.b64encode(file_stream.getvalue()).decode('utf-8')
            
            # Формируем ответ с обоими форматами
            response_data = {
                "success": True,
                "table_data": table_data,
                "file": {
                    "filename": "detectors_table.docx",
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "base64": file_base64,
                    "size": len(file_base64)  # Размер в байтах после кодирования
                },
                "detectors_count": len(table_data),
                "expanded_detectors": expanded_detectors
            }
            
            # Устанавливаем заголовок для скачивания файла
            response = Response(response_data, status=status.HTTP_200_OK)
            response['Content-Disposition'] = 'attachment; filename=detectors_table.docx'
            
            return response

        except Exception as e:
            return Response(
                {"error": f"Ошибка при генерации таблицы: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _create_document(self, table_data):
        """Создает DOCX документ из данных таблицы"""
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        section = document.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(1.8 / 2.54)
        section.left_margin = Inches(0.71 / 2.54)
        section.right_margin = Inches(1.02 / 2.54) 
        section.bottom_margin = Inches(1.55 / 2.54)
        section.gutter = Inches(0)

        table = document.add_table(rows=1, cols=10)
        table.style = 'Table Grid'

        headers = [
            'Детектор', '№ Входа платы IO/входа ДК', 'КИ ПД-2', 'КИ ПД-16',
            '№ Направления', 'Вынос, м', 'Запрос', 'Разрыв, с',
            'Авария незанят, мин', 'Авария занят, мин'
        ]

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0, 31, 95)
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        # Добавляем строки с данными
        for row_data in table_data:
            row_cells = table.add_row().cells

            for i, key in enumerate([
                'name', 'io_board_input', 'ki_pd_2', 'ki_pd_16', 
                'direction_number', 'offset', 'request', 'gap', 
                'unoccupied_alarm', 'occupied_alarm'
            ]):
                row_cells[i].text = str(row_data.get(key, ''))
                cell = row_cells[i]
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

        return document

    def _expand_detectors(self, detectors):
        expanded = []
        for detector in detectors:
            # Проверяем наличие суффикса T
            has_t_suffix = detector.endswith('T')
            base_detector = detector[:-1] if has_t_suffix else detector
            
            if base_detector.startswith('TVP'):
                prefix = 'ТВП '
                parts = base_detector[3:].split('.')
            elif base_detector.startswith('DO'):
                prefix = 'DO'
                parts = base_detector[2:].split('.')
            elif base_detector.startswith('D'):
                prefix = 'D'
                parts = base_detector[1:].split('.')
            else:
                expanded.append(detector)
                continue

            if len(parts) != 2:
                expanded.append(detector)
                continue

            try:
                direction, number = int(parts[0]), int(parts[1])
                for n in range(1, number + 1):
                    expanded_detector = f"{prefix}{direction}.{n}"
                    # Добавляем суффикс T обратно, если он был
                    if has_t_suffix:
                        expanded_detector += 'T'
                    if expanded_detector not in expanded:
                        expanded.append(expanded_detector)
            except ValueError:
                expanded.append(detector)

        return expanded

    def _calculate_ki_pd16_map(self, detectors):
        result = {}
        # Фильтрация детекторов, исключая ТВП
        filtered_detectors = [name for name in detectors if not name.startswith('ТВП')]
        
        # Получаем все данные детекторов для проверки выноса
        all_detectors_data = {}
        for idx, name in enumerate(filtered_detectors, start=1):
            data = self._generate_detector_data(name, idx, detectors, {})
            all_detectors_data[name] = data
        
        board_number = 1
        input_number = 1
        
        for i, name in enumerate(filtered_detectors):
            # Формируем строку КИ ПД-16 в формате 'плата.вход'
            result[name] = f"{board_number}.{input_number}"
            
            # Получаем значение выноса для текущего детектора
            current_offset = all_detectors_data[name]['offset']
            
            # Получаем значение выноса для следующего детектора (если существует)
            next_offset = None
            if i + 1 < len(filtered_detectors):
                next_detector_name = filtered_detectors[i + 1]
                next_offset = all_detectors_data[next_detector_name]['offset']
            
            # Проверяем условия для перехода на следующую плату
            if input_number == 14:
                if next_offset in ["2"]:
                    board_number += 1
                    input_number = 1
                else:
                    input_number += 1
            elif input_number == 15:
                if current_offset in ["7", "10"] or (next_offset in ["2"]):
                    board_number += 1
                    input_number = 1
                else:
                    input_number += 1
            elif input_number == 16:
                if current_offset in ["7", "10"]:
                    board_number += 1
                    input_number = 1
                else:
                    input_number += 1
            else:
                input_number += 1
            
            # Ограничиваем максимальный номер входа 16
            if input_number > 16:
                board_number += 1
                input_number = 1
        
        return result

    def _generate_detector_data(self, detector_name, position, all_detectors, ki_pd16_map):
        try:
            # Определяем тип детектора и наличие суффикса T
            has_t_suffix = detector_name.endswith('T')
            base_name = detector_name[:-1] if has_t_suffix else detector_name
            
            if base_name.startswith('ТВП'):
                detector_type = 'ТВП'
                direction = int(base_name[4:].split('.')[0])
            elif base_name.startswith('DO'):
                detector_type = 'DO'
                direction = int(base_name[2:].split('.')[0])
            elif base_name.startswith('D'):
                detector_type = 'D'
                direction = int(base_name[1:].split('.')[0])
            else:
                detector_type = 'UNKNOWN'
                direction = 0

            # Для детекторов с суффиксом T
            if has_t_suffix:
                if detector_type == 'D':
                    data = {
                        'name': detector_name,
                        'io_board_input': self._generate_io_input(position, ki_pd16_map),
                        'ki_pd_2': '-',
                        'ki_pd_16': ki_pd16_map.get(detector_name, '-'),
                        'direction_number': direction,
                        'offset': '-',  # прочерк для выноса
                        'request': '-',
                        'gap': '-',  # прочерк для разрыва
                        'unoccupied_alarm': '420',  # 420 минут для аварии незанят
                        'occupied_alarm': '3',  # 3 минуты для аварии занят
                    }
                elif detector_type == 'DO':
                    data = {
                        'name': detector_name,
                        'io_board_input': self._generate_io_input(position, ki_pd16_map),
                        'ki_pd_2': '-',
                        'ki_pd_16': ki_pd16_map.get(detector_name, '-'),
                        'direction_number': direction,
                        'offset': '-',  # прочерк для выноса
                        'request': '-',
                        'gap': '-',  # прочерк для разрыва
                        'unoccupied_alarm': '-',  # прочерк для аварии незанят
                        'occupied_alarm': '-',  # прочерк для аварии занят
                    }
                else:
                    # Для других типов с суффиксом T используем стандартную логику
                    data = self._create_standard_data(detector_name, detector_type, direction, position, all_detectors, ki_pd16_map)
            else:
                # Стандартная логика для детекторов без суффикса T
                data = self._create_standard_data(detector_name, detector_type, direction, position, all_detectors, ki_pd16_map)

            return data
        except Exception as e:
            return {'name': detector_name, 'error': f"Ошибка обработки детектора: {str(e)}"}

    def _create_standard_data(self, detector_name, detector_type, direction, position, all_detectors, ki_pd16_map):
        """Создает стандартные данные для детектора без суффикса T"""
        return {
            'name': detector_name,
            'io_board_input': self._generate_io_input(position, ki_pd16_map),
            'ki_pd_2': '' if detector_type == 'ТВП' else '-',
            'ki_pd_16': '' if detector_type == 'ТВП' else ki_pd16_map.get(detector_name, '-'),
            'direction_number': direction,
            'offset': self._generate_offset(detector_type, direction, detector_name, all_detectors),
            'request': '+' if detector_type == 'ТВП' else '-',
            'gap': self._generate_gap(detector_type),
            'unoccupied_alarm': self._generate_unoccupied_alarm(detector_type),
            'occupied_alarm': self._generate_occupied_alarm(detector_type),
        }

    def _generate_io_input(self, position, ki_pd16_map):
        # Определяем количество КИ ПД-16 (первая цифра в последней строке)
        last_ki_pd16 = None
        for detector_name in reversed(ki_pd16_map.keys()):
            if ki_pd16_map[detector_name]:
                last_ki_pd16 = ki_pd16_map[detector_name]
                break
        
        ki_pd16_count = 0
        if last_ki_pd16:
            try:
                ki_pd16_count = int(last_ki_pd16.split('.')[0])
            except:
                pass
        
        # Для первой платы ДК учитываем количество КИ ПД-16
        if position <= (32 - ki_pd16_count):
            board_number = 1
            input_number = position
        else:
            # Для остальных плат стандартный расчет
            adjusted_position = position + ki_pd16_count
            board_number = (adjusted_position - 1) // 32 + 1
            input_number = (adjusted_position - 1) % 32 + 1
        
        return f"{board_number}.{input_number}"

    def _generate_offset(self, detector_type, direction, detector_name, all_detectors):
        if detector_type != 'D':
            return '-'
        try:
            detectors_in_direction = [
                d for d in all_detectors if d.startswith(f'D{direction}') and not d.startswith(f'DO{direction}')
            ]
            count = len(detectors_in_direction)
            
            # Логика для деления по 2, если количество детекторов делится на 2 и не более 4
            if count % 2 == 0 and count <= 4:
                pattern = ['3', '7']  # 7 будет заменено на 6 позже, 3 на 2
            else:
                pattern = ['2', '6', '10']

            idx = detectors_in_direction.index(detector_name) % len(pattern)
            return pattern[idx]
        except:
            return '2'

    def _generate_gap(self, detector_type):
        return '3' if detector_type == 'D' else '-'

    def _generate_unoccupied_alarm(self, detector_type):
        return {'ТВП': '1440', 'D': '180'}.get(detector_type, '-')

    def _generate_occupied_alarm(self, detector_type):
        return {'ТВП': '10', 'D': '30'}.get(detector_type, '-')
